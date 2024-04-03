from flask import Flask, jsonify, request, send_file
import markdown
from docx import Document
from docx.shared import Inches
from io import BytesIO
import requests
import json
import os
from flask_cors import CORS
from bs4 import BeautifulSoup
import aspose.words as aw
from Markdown2docx import Markdown2docx



app = Flask(__name__)
CORS(app)


# Initializing the App and Gemini API
working_dir = os.path.dirname(os.path.abspath(__file__))
config_file_path = f"{working_dir}/config.json"
config_data = json.load(open(config_file_path))
GOOGLE_API_KEY = config_data["GOOGLE_API_KEY"]

class RolePlayCreator:
    def __init__(self, api_key):
        self.api_key = api_key
        self.base_url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-pro:generateContent?key=" + self.api_key

    @staticmethod
    def construct_prompt(data):
        prompt_parts = [
            "Generate a role-play exercise with the following specifications:\n",
            f"**Industry:** {data['industry']}\n",
            f"**Learning Objective:** {data['learningObjective']}\n",
            "**Roles:**\n"
        ]
        prompt_parts.extend([f"- {role}\n" for role in data['roles']])
        prompt_parts.append("**Required Skills:**\n")
        prompt_parts.extend([f"- {skill}\n" for skill in data['skills']])

        if 'complexity' in data:
            prompt_parts.append(f"**Complexity:** {data['complexity']}\n")

        if 'experience' in data:
            prompt_parts.append(f"**Experience Level:** {data['experience']}\n")

        if 'industryContext' in data:
            prompt_parts.append(f"**Industry Context:** {data['industryContext']}\n")

        if 'scenarioSettings' in data:
            prompt_parts.append(f"**Scenario Settings:** {data['scenarioSettings']}\n")

        return ''.join(prompt_parts)

rolePlayCreator_service = RolePlayCreator(api_key=GOOGLE_API_KEY)

@app.route('/role_play_creator', methods=['POST'])
def rolePlayCreator():

    try:
        headers = {'Content-Type': 'application/json'}
        request_data = request.json
        prompt = rolePlayCreator_service.construct_prompt(request_data)
        generation_config = {
            'temperature': 0.9,
            'topK': 1,
            'topP': 1,
            'maxOutputTokens': 2048,
            'stopSequences': []
        }

        request_body = {
            'contents': [{'parts': [{'text': prompt}]}],
            'generationConfig': generation_config
        }

        response = requests.post(rolePlayCreator_service.base_url, json=request_body, headers=headers)
        response.raise_for_status()
        response_data = response.json()
        rolePlayExercie = [candidate['content']['parts'][0]['text'] for candidate in response_data['candidates']]

        return jsonify(rolePlayExercie)

    except Exception as e:
        print("Service Exception:", str(e))
        raise Exception("Error in getting response from Gemini API")

@app.route('/download-docx', methods=['POST'])
def download_docx():
    markdown_content = request.json['markdown_content']
    file_path = working_dir+"/RolePlay.md"

    create_md_file(markdown_content, file_path)
    output = aw.Document()
    output.remove_all_children()
    project = Markdown2docx(working_dir+"/RolePlay")
    project.eat_soup()
    project.save()
    # doc = aw.Document(working_dir+"/example.md")
    # doc.save("Output.docx")
    return send_file("RolePlay.docx", as_attachment=True, download_name="role-play.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    # input = aw.Document(working_dir+"/example.md")
    # output.append_document(markdown_content, aw.ImportFormatMode.KEEP_SOURCE_FORMATTING)
    # output.save("Output.docx")




    try:
        markdown_content = request.json['markdown_content']


        # Convert Markdown content to HTML
        html_content = markdown.markdown(markdown_content)

        # Parse HTML content
        soup = BeautifulSoup(html_content, 'html.parser')

        # Create a new Document
        doc = Document()

        # Process each tag in the HTML content
        for tag in soup.find_all(True):
            if tag.name == 'h1':
                # Handle heading level 1
                doc.add_heading(tag.text, level=1)
            elif tag.name == 'h2':
                # Handle heading level 2
                doc.add_heading(tag.text, level=2)
            elif tag.name == 'h3':
                # Handle heading level 3
                doc.add_heading(tag.text, level=3)
            elif tag.name == 'p':
                # Handle paragraph
                doc.add_paragraph(tag.text)
            elif tag.name == 'ul':
                # Handle unordered list
                for li in tag.find_all('li'):
                    doc.add_paragraph(li.text, style='ListBullet')
            elif tag.name == 'ol':
                # Handle ordered list
                for li in tag.find_all('li'):
                    doc.add_paragraph(li.text, style='ListNumber')

        # Create a buffer to store the document
        doc_buffer = BytesIO()

        # Save the document to the buffer
        doc.save(doc_buffer)

        # Move to the beginning of the buffer
        doc_buffer.seek(0)

        # Send the buffer as a downloadable file
        return send_file(doc_buffer, as_attachment=True, download_name="role-play.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        print("Error:", str(e))
        return jsonify({'error': 'Failed to generate DOCX file.'}), 500
    

    
def create_md_file(text_content, file_path):
    try:
        # Open the file in write mode
        with open(file_path, 'w') as f:
            # Write the text content to the file
            f.write(text_content)
        print(f"Markdown file '{file_path}' created successfully.")
    except Exception as e:
        print("Error:", str(e))    
        
# if __name__ == '__main__':
#     app.run(debug=True, host="0.0.0.0")



